#!/usr/bin/env python3
"""
replace_chinese.py

Replace Chinese words in .txt or .docx files using mappings (Chinese=English per line).
Features:
 - keep first mapping when duplicates exist
 - prefer longest match when multiple keys overlap
 - add a space before and after the replacement English text
 - preserve docx run formatting as far as possible
 - --fast-serialize: trade some formatting accuracy for much faster paragraph reading on large .docx
 - progress bars (tqdm) shown during serialization and replacement
 - parallel processing (multiprocessing)
 - prints total replacements and writes output file
"""
from __future__ import annotations
import argparse
import re
import sys
import os
from collections import OrderedDict
from typing import List, Dict, Tuple, Any
from multiprocessing import Pool, cpu_count
from tqdm import tqdm

# docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except Exception:
    Document = None

# -------------------------
# Mapping utilities
# -------------------------
def load_mappings(mapping_path: str) -> OrderedDict:
    mapping = OrderedDict()
    with open(mapping_path, 'r', encoding='utf-8-sig') as f:
        for raw in f:
            line = raw.strip()
            if not line or line.startswith('#'):
                continue
            if '=' not in line:
                continue
            key, val = line.split('=', 1)
            key = key.strip()
            val = val.strip()
            if key and key not in mapping:
                mapping[key] = val
    return mapping

def build_pattern_string_from_keys(keys: List[str]) -> str:
    # sort by length descending to favor longest matches
    keys_sorted = sorted(keys, key=lambda s: len(s), reverse=True)
    escaped = [re.escape(k) for k in keys_sorted]
    pattern_str = '(' + '|'.join(escaped) + ')'
    return pattern_str

# -------------------------
# TXT processing (line-level, parallel)
# -------------------------
def replace_in_line_worker(args: Tuple[str, str, Dict[str,str]]):
    """
    Worker for txt: receives (line, pattern_str, mapping)
    Recompile pattern in worker for reliability across platforms.
    Returns (new_line, count)
    """
    line, pattern_str, mapping = args
    pattern = re.compile(pattern_str)
    count = 0
    def repl(m):
        nonlocal count
        key = m.group(0)
        repl_text = mapping.get(key, key)
        count += 1
        return ' ' + repl_text + ' '
    new_line = pattern.sub(repl, line)
    return new_line, count

def process_txt(input_path: str, output_path: str, pattern_str: str, mapping: Dict[str,str], workers:int) -> int:
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    total_lines = len(lines)
    args_iter = ((lines[i], pattern_str, mapping) for i in range(total_lines))

    replaced_lines = [None] * total_lines
    total_replacements = 0

    # Use imap to maintain order and show progress
    with Pool(processes=workers) as p:
        for i, (new_line, count) in enumerate(tqdm(p.imap(replace_in_line_worker, args_iter, chunksize=200),
                                                   total=total_lines, desc="Processing lines")):
            replaced_lines[i] = new_line
            total_replacements += count

    with open(output_path, 'w', encoding='utf-8') as f:
        f.writelines(replaced_lines)

    return total_replacements

# -------------------------
# DOCX helpers: serialization & rebuild
# -------------------------
def paragraph_to_serial(paragraph) -> Dict[str, Any]:
    """
    Full serialization: capture run.text + common style attributes per run.
    Returns {'runs': [{'text':..., 'bold':..., 'italic':..., 'underline':..., 'style':..., 'font_name':..., 'font_size':..., 'font_color':...}, ...], 'text': concatenated}
    """
    runs_info = []
    for run in paragraph.runs:
        font = run.font
        color_rgb = None
        try:
            if font.color and font.color.rgb:
                color_rgb = str(font.color.rgb)
        except Exception:
            color_rgb = None
        runs_info.append({
            'text': run.text or '',
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'style': run.style.name if run.style is not None else None,
            'font_name': font.name,
            'font_size': font.size.pt if font.size is not None else None,
            'font_color': color_rgb,
        })
    full_text = ''.join(r['text'] for r in runs_info)
    return {'runs': runs_info, 'text': full_text}

def paragraph_to_serial_fast(paragraph) -> Dict[str, Any]:
    """
    Fast serialization: collect only run.text (and minimal style anchor: bools). Much faster for huge docs.
    """
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text or '',
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            # don't collect font sizes/colors/names to save time
            'style': run.style.name if run.style is not None else None,
        })
    full_text = ''.join(r['text'] for r in runs_info)
    return {'runs': runs_info, 'text': full_text}

def process_paragraph_serial_worker(args: Tuple[int, Dict[str,Any], str, Dict[str,str]]):
    """
    Worker for paragraphs. Receives (idx, serial_dict, pattern_str, mapping).
    Recompile regex and perform replacements on serial_dict['text'].
    Return (idx, rebuilt_chunks_list, replacements_count)
    rebuilt_chunks_list = [(text, style_dict), ...] where style_dict contains the serialization's style fields.
    """
    idx, serial, pattern_str, mapping = args
    pattern = re.compile(pattern_str)
    full_text: str = serial['text']
    runs_info = serial['runs']

    if not full_text:
        return idx, [], 0

    # Build per-character style mapping
    char_styles = []
    for run_idx, r in enumerate(runs_info):
        style = {
            # copy whatever style keys exist in run info
            'bold': r.get('bold'),
            'italic': r.get('italic'),
            'underline': r.get('underline'),
            'style': r.get('style'),
            'font_name': r.get('font_name') if 'font_name' in r else None,
            'font_size': r.get('font_size') if 'font_size' in r else None,
            'font_color': r.get('font_color') if 'font_color' in r else None,
        }
        for _ in r['text']:
            char_styles.append(style)

    def style_for_index(i: int) -> Dict[str,Any]:
        if 0 <= i < len(char_styles):
            return char_styles[i]
        else:
            # fallback
            return {'bold': None, 'italic': None, 'underline': None, 'style': None, 'font_name': None, 'font_size': None, 'font_color': None}

    matches = list(pattern.finditer(full_text))
    if not matches:
        # no change: return original runs as chunks
        rebuilt = []
        for r in runs_info:
            rebuilt.append((r['text'], {
                'bold': r.get('bold'),
                'italic': r.get('italic'),
                'underline': r.get('underline'),
                'style': r.get('style'),
                'font_name': r.get('font_name') if 'font_name' in r else None,
                'font_size': r.get('font_size') if 'font_size' in r else None,
                'font_color': r.get('font_color') if 'font_color' in r else None,
            }))
        return idx, rebuilt, 0

    rebuilt_chunks = []
    pos = 0
    replacements_count = 0

    for m in matches:
        s, e = m.start(), m.end()
        if s > pos:
            before_text = full_text[pos:s]
            # split by style changes
            run_start = 0
            while run_start < len(before_text):
                abs_i = pos + run_start
                cur_style = style_for_index(abs_i)
                seg_len = 1
                while run_start + seg_len < len(before_text):
                    if style_for_index(pos + run_start + seg_len) == cur_style:
                        seg_len += 1
                    else:
                        break
                rebuilt_chunks.append((before_text[run_start:run_start+seg_len], cur_style))
                run_start += seg_len

        key = m.group(0)
        repl_text = ' ' + mapping.get(key, key) + ' '
        repl_style = style_for_index(s)
        rebuilt_chunks.append((repl_text, repl_style))
        replacements_count += 1
        pos = e

    # tail
    if pos < len(full_text):
        tail_text = full_text[pos:]
        run_start = 0
        while run_start < len(tail_text):
            abs_i = pos + run_start
            cur_style = style_for_index(abs_i)
            seg_len = 1
            while run_start + seg_len < len(tail_text):
                if style_for_index(pos + run_start + seg_len) == cur_style:
                    seg_len += 1
                else:
                    break
            rebuilt_chunks.append((tail_text[run_start:run_start+seg_len], cur_style))
            run_start += seg_len

    # merge adjacent chunks with identical style
    merged = []
    for t, s in rebuilt_chunks:
        if not t:
            continue
        if merged and merged[-1][1] == s:
            merged[-1] = (merged[-1][0] + t, s)
        else:
            merged.append((t, s))

    return idx, merged, replacements_count

# -------------------------
# Apply rebuilt runs to docx paragraph
# -------------------------
def clear_paragraph_runs(paragraph):
    # zero out existing runs' text
    for run in paragraph.runs:
        run.text = ''

def _apply_style_to_run(run, style: Dict[str,Any]):
    # best-effort application of style dict fields
    try:
        run.bold = style.get('bold', None)
        run.italic = style.get('italic', None)
        run.underline = style.get('underline', None)
        if style.get('style') is not None:
            try:
                run.style = style['style']
            except Exception:
                pass
        font = run.font
        if style.get('font_name'):
            try:
                font.name = style['font_name']
            except Exception:
                pass
        if style.get('font_size'):
            try:
                font.size = Pt(style['font_size'])
            except Exception:
                pass
        if style.get('font_color'):
            try:
                hexstr = style['font_color']
                if hexstr.startswith('RGBColor('):
                    hexstr = hexstr.split('(')[1].rstrip(')')
                hexstr = hexstr.strip().replace('0x','').upper()
                if len(hexstr) == 6:
                    r = int(hexstr[0:2],16)
                    g = int(hexstr[2:4],16)
                    b = int(hexstr[4:6],16)
                    font.color.rgb = RGBColor(r,g,b)
            except Exception:
                pass
    except Exception:
        pass

def apply_rebuilt_runs_to_paragraph(paragraph, rebuilt_chunks: List[Tuple[str, Dict[str,Any]]]):
    # attempt to reuse existing runs where possible
    existing_runs = paragraph.runs
    # clear existing runs' text
    for r in existing_runs:
        r.text = ''
    idx = 0
    for chunk_text, style in rebuilt_chunks:
        if idx < len(existing_runs):
            r = existing_runs[idx]
            r.text = chunk_text
            _apply_style_to_run(r, style)
        else:
            r = paragraph.add_run(chunk_text)
            _apply_style_to_run(r, style)
        idx += 1

# -------------------------
# DOCX processing orchestration
# -------------------------
def process_docx(input_path: str, output_path: str, pattern_str: str, mapping: Dict[str,str], workers:int, fast_serialize: bool) -> int:
    if Document is None:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")

    print("Opening document and starting to serialize paragraphs...")  # immediate feedback
    doc = Document(input_path)

    all_paragraphs = []
    # collect top-level paragraphs
    for paragraph in doc.paragraphs:
        all_paragraphs.append(paragraph)
    # collect paragraphs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_paragraphs.append(paragraph)

    total_paras = len(all_paragraphs)
    if total_paras == 0:
        doc.save(output_path)
        return 0

    # serialize with progress bar (fast or full)
    serials = []
    desc = "Serializing paragraphs (fast)" if fast_serialize else "Serializing paragraphs (full)"
    serializer = paragraph_to_serial_fast if fast_serialize else paragraph_to_serial
    for p in tqdm(all_paragraphs, desc=desc, unit="para"):
        serials.append(serializer(p))

    # prepare args for workers
    args = [(idx, serials[idx], pattern_str, mapping) for idx in range(len(serials))]
    results = [None] * len(args)
    total_replacements = 0

    # choose chunksize heuristically
    chunksize = max(1, min(200, len(args) // (workers * 2 + 1)))

    with Pool(processes=workers) as pool:
        for idx, rebuilt, count in tqdm(pool.imap_unordered(process_paragraph_serial_worker, args, chunksize=chunksize),
                                        total=len(args),
                                        desc="Processing paragraphs (replacing)"):
            results[idx] = (rebuilt, count)
            total_replacements += count

    # apply rebuilt runs back into doc (main process)
    for paragraph, (rebuilt, _) in zip(all_paragraphs, results):
        if rebuilt is None:
            continue
        clear_paragraph_runs(paragraph)
        apply_rebuilt_runs_to_paragraph(paragraph, rebuilt)

    doc.save(output_path)
    return total_replacements

# -------------------------
# Main
# -------------------------
if __name__ == "__main__":
    # File with Chinese=English lines
    # mapping_file = "C:\\DATA\\Novels\\Pokemon_Glossary_Word_Replacement.txt"
    mapping_file = "C:\\DATA\\Novels\\Conan_Glossary_Word_Replacement.txt"
    # mapping_file = "C:\\DATA\\Novels\\HarryPotter_Glossary_Word_Replacement.txt"
    # mapping_file = "C:\\DATA\\Novels\\OnePiece_Glossary_Word_Replacement.txt"
    # mapping_file = "C:\\DATA\\Novels\\GenshinHonkai_Glossary_Word_Replacement.txt"
    # mapping_file = "C:\\DATA\\Novels\\FairyTail_Glossary_Word_Replacement.txt"
    
    file_root = "C:\\DATA\\Novels\\Isn't it very scientific to kill people with the Death Note"
    
    input_file = os.path.join(file_root, "input.txt")  # can be input.txt or input.docx
    output_file = os.path.join(file_root, "output.txt")  # .txt or .docx
    fast_serialize = True
    workers = max(1, cpu_count()-1)

    mapping = load_mappings(mapping_file)
    if not mapping:
        print("No mappings loaded from", mapping_file)
        sys.exit(1)

    pattern = build_pattern_string_from_keys(list(mapping.keys()))
    ext = os.path.splitext(input_file)[1].lower()

    print(f"Loaded {len(mapping)} mappings. Using {workers} worker(s).")
    total_replacements = 0

    if ext == '.txt':
        total_replacements = process_txt(input_file, output_file, pattern, mapping, workers)
    elif ext == '.docx':
        total_replacements = process_docx(input_file, output_file, pattern, mapping, workers, fast_serialize)
    else:
        print("Unsupported input file type. Only .txt and .docx supported.")
        sys.exit(1)

    print(f"Done. Total replacements made: {total_replacements}")
    print(f"Output written to: {output_file}")