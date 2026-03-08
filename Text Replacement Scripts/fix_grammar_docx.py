import os
import sys
import docx
import language_tool_python
from tqdm import tqdm

def process_docx_fast(input_file, output_file):
    """
    Fast, free grammar correction using LanguageTool and multi-stage progress bars.
    """
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        return

    print("Initializing Grammar Engine (this happens once)...")
    # 'en-US' is standard. Change to 'en-GB' for British English.
    tool = language_tool_python.LanguageTool('en-US') 

    print(f"Loading {input_file}...")
    try:
        doc = docx.Document(input_file)
    except Exception as e:
        print(f"Failed to open document: {e}")
        return

    # --- Pre-calculation for Progress Bars ---
    total_paragraphs = len(doc.paragraphs)
    
    # Calculate total table cells
    total_table_tasks = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_table_tasks += len(cell.paragraphs)

    print(f"\nProcessing document: {input_file}")
    print("---------------------------------------------------")

    # --- Phase 1: Main Document Body ---
    # We use 'unit_scale=True' and a simple format for speed
    with tqdm(total=total_paragraphs, desc="Main Body", colour="green", unit="para") as pbar:
        for para in doc.paragraphs:
            # OPTIMIZATION: Skip empty strings immediately to save processing time
            if len(para.text.strip()) > 0:
                try:
                    # Apply correction
                    para.text = tool.correct(para.text)
                except Exception:
                    # If a specific paragraph fails, skip it so the script doesn't crash
                    pass
            pbar.update(1)

    # --- Phase 2: Tables ---
    # Only show this bar if tables actually exist
    if total_table_tasks > 0:
        print("\n") # formatting spacer
        with tqdm(total=total_table_tasks, desc="Tables   ", colour="cyan", unit="cell") as pbar_tables:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if len(para.text.strip()) > 0:
                                try:
                                    para.text = tool.correct(para.text)
                                except Exception:
                                    pass
                            pbar_tables.update(1)
    else:
        print("\nNo tables found. Skipping Phase 2.")

    # --- Phase 3: Save ---
    print("\nSaving file...")
    try:
        doc.save(output_file)
        print(f"Success! ✅ File saved as: {output_file}")
    except PermissionError:
        print("❌ Error: Could not save file. Please close the Word document if it is open.")

if __name__ == "__main__":
    # You can hardcode the filename here
    file_root = "C:\\DATA\\Novels\\Pokémon Adventures - GAMER PATH"
    input_file = os.path.join(file_root, "input.docx")
    output_file = os.path.join(file_root, "output.docx")
        
    process_docx_fast(input_file, output_file)