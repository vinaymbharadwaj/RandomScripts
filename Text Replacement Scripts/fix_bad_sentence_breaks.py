import os
import re
from pathlib import Path
from tqdm import tqdm

try:
    import docx
except ImportError:
    print("Please install required libraries: pip install python-docx tqdm")
    exit()

def copy_run_format(old_run, new_run):
    """Copies all font styles from one run to another."""
    new_run.bold = old_run.bold
    new_run.italic = old_run.italic
    new_run.underline = old_run.underline
    new_run.font.strike = old_run.font.strike
    new_run.font.superscript = old_run.font.superscript
    new_run.font.subscript = old_run.font.subscript
    
    if old_run.font.name:
        new_run.font.name = old_run.font.name
    if old_run.font.size:
        new_run.font.size = old_run.font.size
    if old_run.font.color and old_run.font.color.rgb:
        new_run.font.color.rgb = old_run.font.color.rgb

def reformat_novel_docx(input_path, output_path, sentences_per_paragraph=4):
    print(f"\nLoading '{input_path.name}'... (This might take a moment)")
    doc = docx.Document(input_path)
    new_doc = docx.Document()
    
    paragraph_buffer = []
    
    # Regex to count actual sentence endings (ignoring titles like Mr. and Mrs.)
    sentence_end_pattern = re.compile(r'(?<!\bMt)(?<!\bMr)(?<!\bMrs)(?<!\bMs)(?<!\bDr)(?<!\bProf)([.!?][\"\']?)')

    def flush_buffer():
        """Merges buffered short lines into one formatted paragraph."""
        if not paragraph_buffer:
            return
            
        new_para = new_doc.add_paragraph()
        
        # Inherit the paragraph style and alignment from the first line in the buffer
        new_para.style = paragraph_buffer[0].style
        new_para.alignment = paragraph_buffer[0].alignment
        
        for i, old_para in enumerate(paragraph_buffer):
            # SMART SPACING: Decide if we need a space between merged lines
            if i > 0:
                prev_text = new_para.text
                next_text = old_para.text.lstrip()
                if prev_text and not prev_text.endswith((' ', '-', '—')):
                    if next_text and not next_text.startswith((',', '.', '!', '?', ':', ';')):
                        new_para.add_run(' ')
            
            for run in old_para.runs:
                # Normalizing text and handling line breaks
                text = run.text.replace('\n', ' ')
                
                # ELLIPSES CLEANUP: Normalize messy dots into a clean '...'
                text = re.sub(r'\.{2,}|…+', '...', text)
                text = re.sub(r',{2,}', ',', text)
                
                if text:
                    new_run = new_para.add_run(text)
                    copy_run_format(run, new_run)
                    
        paragraph_buffer.clear()

    # Process paragraphs with a progress bar
    for para in tqdm(doc.paragraphs, desc="Rebuilding Paragraphs", unit="par"):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        
        # 1. PRESERVE HEADINGS & TITLES
        if 'Heading' in style_name or 'Title' in style_name or 'Subtitle' in style_name:
            flush_buffer() # Empty out normal text first
            
            new_heading = new_doc.add_paragraph()
            new_heading.style = para.style
            new_heading.alignment = para.alignment
            
            for run in para.runs:
                run_text = run.text.replace('\n', ' ')
                
                # Clean up ellipses inside headings too
                run_text = re.sub(r'\.{2,}|…+', '...', run_text)
                run_text = re.sub(r',{2,}', ',', run_text)
                
                if run_text:
                    new_run = new_heading.add_run(run_text)
                    copy_run_format(run, new_run)
            continue
            
        # 2. SKIP EMPTY LINES
        if not text:
            continue
            
        # 3. ISOLATE AUTHOR/TRANSLATOR NOTES IN BRACKETS
        is_bracketed = (text.startswith('(') and text.endswith(')')) or \
                       (text.startswith('[') and text.endswith(']')) or \
                       (text.startswith('（') and text.endswith('）')) or \
                       (text.startswith('【') and text.endswith('】'))
                       
        if is_bracketed:
            flush_buffer()            # Print whatever normal text we were holding onto
            paragraph_buffer.append(para) 
            flush_buffer()            # Instantly print the bracketed line so it stays isolated
            continue
            
        # 4. BUFFER NORMAL TEXT
        paragraph_buffer.append(para)
        
        # Combine the text in the buffer to check how many sentences we've accumulated
        combined_text = " ".join([p.text for p in paragraph_buffer])
        sentence_count = len(sentence_end_pattern.findall(combined_text))
        
        # Flush if we hit our sentence limit OR if the text ends with dialogue
        if sentence_count >= sentences_per_paragraph or text.endswith(('"', '”', "'")):
            flush_buffer()
            
    # Flush anything remaining at the very end
    flush_buffer()
    
    print(f"Saving new document...")
    new_doc.save(output_path)
    print(f"Done! Saved as: {output_path.name}\n")

def main():
    # --- SETUP FILE PATHS ---
    # --- ENTER YOUR FILE PATH HERE ---
    FILE_ROOT = "C:\\DATA\\Novels\\Eat Reba's soft rice, spend all her small treasury at the beginning"
    FILE_NAME = os.path.join(FILE_ROOT, "input.docx")
    input_file = Path(FILE_NAME)
    
    if not input_file.exists():
        print(f"Error: Could not find '{input_file.name}'. Check the path!")
        return
        
    output_file = input_file.parent / f"cleaned_{input_file.name}"
    
    # Run the formatter (groups 4 sentences per paragraph)
    reformat_novel_docx(input_file, output_file, sentences_per_paragraph=4)

if __name__ == '__main__':
    main()