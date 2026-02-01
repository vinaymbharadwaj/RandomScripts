import os
from docx import Document
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm

def fix_chinese_formatting_worker(file_info):
    """
    Worker function to process a single file. 
    Preserves Headings 1-9 while fixing body text line breaks.
    """
    input_path, output_path = file_info
    
    try:
        doc = Document(input_path)
        # We create a new document based on the default template, 
        # which contains standard styles like 'Heading 1', 'Normal', etc.
        new_doc = Document()
        
        paragraph_buffer = ""
        # Punctuation that marks the end of a sentence
        terminal_punctuation = tuple(['。', '！', '？', '…', '”', '.', '!', '?', '"'])

        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name
            
            # Skip empty lines
            if not text:
                continue

            # --- CHECK FOR HEADINGS ---
            # If the style is a Heading, we treat it as a "Hard Break"
            if style_name.startswith("Heading"):
                # 1. Flush any body text currently in the buffer
                if paragraph_buffer:
                    new_doc.add_paragraph(paragraph_buffer)
                    paragraph_buffer = "" 
                
                # 2. Add the heading immediately
                new_para = new_doc.add_paragraph(text)
                
                # 3. Apply the style (Try/Except in case style doesn't exist in new doc)
                try:
                    new_para.style = style_name
                except KeyError:
                    # Fallback: If 'Heading 1' isn't found, just make it bold
                    run = new_para.runs[0]
                    run.bold = True
                
                continue # Move to next paragraph, do not process this as body text

            # --- PROCESS BODY TEXT (Merge Logic) ---
            paragraph_buffer += text
            
            if text.endswith(terminal_punctuation):
                new_doc.add_paragraph(paragraph_buffer)
                paragraph_buffer = "" 
            else:
                pass 

        # Flush whatever is left in the buffer at the end of the doc
        if paragraph_buffer:
            new_doc.add_paragraph(paragraph_buffer)

        new_doc.save(output_path)
        return (True, input_path)

    except Exception as e:
        return (False, f"{input_path} Error: {e}")

def process_folder_parallel(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    tasks = []
    # Scan for files
    files = [f for f in os.listdir(input_folder) if f.endswith(".docx") and not f.startswith("~$")]
    
    if not files:
        print(f"No .docx files found in {input_folder}")
        return

    for filename in files:
        in_path = os.path.join(input_folder, filename)
        out_path = os.path.join(output_folder, f"Fixed_{filename}")
        tasks.append((in_path, out_path))

    print(f"Starting parallel processing on {len(tasks)} files...")

    with ProcessPoolExecutor() as executor:
        futures = [executor.submit(fix_chinese_formatting_worker, task) for task in tasks]
        
        with tqdm(total=len(tasks), unit="file", desc="Fixing Docs") as pbar:
            for future in as_completed(futures):
                success, message = future.result()
                if not success:
                    tqdm.write(f"FAILED: {message}")
                pbar.update(1)

    print(f"\nDone! Files are in '{output_folder}'")


# --- Usage ---
if __name__ == "__main__":
    # You can change these folder names to whatever you like
    file_root = "C:\\DATA\\Novels\\Isn't it very scientific to kill people with the Death Note"
    
    target_input_folder = os.path.join(file_root, "output_parts") 
    target_output_folder = os.path.join(file_root, "result_files")

    process_folder_parallel(target_input_folder, target_output_folder)