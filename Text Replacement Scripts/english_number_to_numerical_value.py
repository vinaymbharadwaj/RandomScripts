import re
import sys
import os

try:
    from docx import Document
except ImportError:
    print("Warning: 'python-docx' library is not installed.")
    print("You will only be able to process .txt files.")
    print("To process .docx files, run: pip install python-docx")


# --- Progress Bar Utility ---

def print_progress_bar(iteration, total, prefix='Progress:', suffix='Complete', decimals=1, length=50, fill='█'):
    """
    Creates a simple console progress bar using the carriage return (\r) for real-time updates.
    """
    if total == 0:
        return
    
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filled_length = int(length * iteration // total)
    bar = fill * filled_length + '-' * (length - filled_length)
    
    # Use sys.stdout.write and flush for real-time update
    sys.stdout.write(f'\r{prefix} |{bar}| {percent}% {suffix}')
    sys.stdout.flush()
    
    # Print New Line on Complete
    if iteration == total: 
        sys.stdout.write('\n')
        sys.stdout.flush()


# --- Number Conversion Logic ---
# This dictionary contains the word-to-number mappings needed for chapter titles
NUM_WORDS = {
    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5, 'six': 6, 'seven': 7, 'eight': 8, 'nine': 9,
    'ten': 10, 'eleven': 11, 'twelve': 12, 'thirteen': 13, 'fourteen': 14, 'fifteen': 15, 'sixteen': 16,
    'seventeen': 17, 'eighteen': 18, 'nineteen': 19, 'twenty': 20, 'thirty': 30, 'forty': 40, 'fifty': 50,
    'sixty': 60, 'seventy': 70, 'eighty': 80, 'ninety': 90
}

def word_to_num(text):
    """
    Converts an English spelled-out number (now including thousands) to its integer equivalent.
    Example: "One Thousand Five Hundred and Thirty Six" -> 1536
    """
    words = text.lower().replace('-', ' ').strip().split()
    
    final_value = 0
    current_value = 0 # The number currently being built (up to the hundreds place)

    for word in words:
        # Ignore common conjunctions
        if word in ['and', 'a']:
            continue
        # Check for simple numbers or tens
        elif word in NUM_WORDS:
            current_value += NUM_WORDS[word]
        # Handle 'hundred'
        elif word == 'hundred':
            # Apply 100 multiplier to the unit built so far (e.g., 'five hundred')
            factor = current_value if current_value > 0 else 1
            current_value = factor * 100
        # Handle 'thousand'
        elif word == 'thousand':
            # Apply 1000 multiplier to the accumulated value and add it to the final result.
            if current_value == 0:
                 current_value = 1 # Implies "one thousand"

            final_value += current_value * 1000
            current_value = 0 # Reset current value to handle remaining digits (e.g., '534' in '2534')
        else:
            # If an unknown word or punctuation remains, it's not a valid number sequence
            return None

    # Add any remaining units/tens/hundreds (the part after 'thousand', or the whole number if no 'thousand' was present)
    final_value += current_value
    
    return final_value if final_value > 0 else None


# --- File Utility Logic ---

def create_output_path(filepath):
    """Creates a new output path with '_edited' appended before the extension."""
    base, ext = os.path.splitext(filepath)
    return f"{base}_edited{ext}"


def process_txt_file(filepath):
    """Reads a TXT file, replaces all chapter numbers, and writes the output to a new file with progress tracking."""
    output_filepath = create_output_path(filepath)
    
    # Helper function for delimited replacements
    def final_replacement_logic(match):
        prefix = match.group(1) # Chapter<space>
        number_words = match.group(2).strip()
        suffix = match.group(3) # Delimiter
        
        num_value = word_to_num(number_words)
        
        if num_value is not None:
            return f"{prefix}{num_value}{suffix}"
        else:
            return match.group(0)

    # Helper function for end-of-line replacements
    def final_replacement_logic_eol(match):
        prefix = match.group(1)
        number_words = match.group(2).strip()
        
        num_value = word_to_num(number_words)
        
        if num_value is not None:
            return f"{prefix}{num_value}"
        else:
            return match.group(0)
        
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        total_lines = len(lines)
        new_lines = []
        
        print_progress_bar(0, total_lines, prefix='TXT Processing:')

        for i, line in enumerate(lines):
            # 1. Handle chapters followed by a delimiter (space, colon, comma, etc.)
            processed_line = re.sub(
                r'(Chapter\s+)([\w\s-]+?)([.,:;!\s])',
                final_replacement_logic,
                line,
                flags=re.IGNORECASE
            )

            # 2. Handle the case where the chapter number is at the end of the content
            processed_line = re.sub(
                r'(Chapter\s+)([\w\s-]+?)\Z',
                final_replacement_logic_eol,
                processed_line,
                flags=re.IGNORECASE
            )
            
            new_lines.append(processed_line)
            print_progress_bar(i + 1, total_lines, prefix='TXT Processing:') # Update progress bar

        # Write the corrected content to the NEW output file
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)

        print(f"Successfully processed TXT file. Output saved to: {output_filepath}")

    except FileNotFoundError:
        print(f"Error: Input file not found at {filepath}")
    except Exception as e:
        print(f"An error occurred while processing the TXT file: {e}")


def process_docx_file(filepath):
    """
    Reads a DOCX file, replaces chapter numbers ONLY in 'Heading 1' styles,
    and saves the modified document to a new file with progress tracking.
    """
    if 'Document' not in globals():
        print(f"Cannot process DOCX file: {filepath}. 'python-docx' library is not available.")
        return

    output_filepath = create_output_path(filepath)

    try:
        document = Document(filepath)
        paragraphs = document.paragraphs
        total_paragraphs = len(paragraphs)
        
        # Helper function for delimited replacements
        regex = r'(Chapter\s+)([\w\s-]+?)([.,:;!\s])'
        
        def final_replacement_logic(match):
            prefix = match.group(1)
            number_words = match.group(2).strip()
            suffix = match.group(3)
            
            num_value = word_to_num(number_words)
            
            if num_value is not None:
                return f"{prefix}{num_value}{suffix}"
            else:
                return match.group(0)
            
        # Helper function for end-of-line replacements
        regex_eol = r'(Chapter\s+)([\w\s-]+?)\Z'
        def final_replacement_logic_eol(match):
            prefix = match.group(1)
            number_words = match.group(2).strip()
            
            num_value = word_to_num(number_words)
            
            if num_value is not None:
                return f"{prefix}{num_value}"
            else:
                return match.group(0)

        
        changes_made = False
        
        print_progress_bar(0, total_paragraphs, prefix='DOCX Processing:')

        for i, paragraph in enumerate(paragraphs):
            # Check for the required style: 'Heading 1'
            if paragraph.style.name.lower() == 'heading 1':
                original_text = paragraph.text
                
                # 1. Apply replacement logic for delimited chapters
                new_text = re.sub(
                    regex, 
                    final_replacement_logic, 
                    original_text, 
                    flags=re.IGNORECASE
                )
                
                # 2. Apply replacement logic for end-of-line chapters
                new_text = re.sub(
                    regex_eol, 
                    final_replacement_logic_eol, 
                    new_text, 
                    flags=re.IGNORECASE
                )


                if new_text != original_text:
                    changes_made = True
                    
                    # Simple approach to update paragraph text content
                    for run in paragraph.runs:
                        run.text = ""
                    
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)
            
            print_progress_bar(i + 1, total_paragraphs, prefix='DOCX Processing:') # Update progress bar


        if changes_made:
            document.save(output_filepath)
            print(f"Successfully processed DOCX file. Output saved to: {output_filepath}")
        else:
            # If no changes are made, save a copy of the original file to the new path
            document.save(output_filepath)
            print(f"No changes made to DOCX file. A copy was saved to: {output_filepath}")


    except FileNotFoundError:
        print(f"Error: Input file not found at {filepath}")
    except Exception as e:
        print(f"An error occurred while processing the DOCX file: {e}")


if __name__ == "__main__":
    file_root = "C:\\DATA\\Novels\\I Repair Cultural Relics for the Country"
    input_file = os.path.join(file_root, "input.docx")  # can be input.txt or input.docx
    
    if input_file.lower().endswith('.txt'):
        process_txt_file(input_file)
    elif input_file.lower().endswith('.docx'):
        process_docx_file(input_file)
    else:
        print("Error: Unsupported file type. Please provide a .txt or .docx file.")
        sys.exit(1)