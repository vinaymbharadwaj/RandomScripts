#!/usr/bin/env python3
"""
Add/replace a colon after leading numbers in "Heading 1" paragraphs of a .docx,
preserving run-level formatting.

Features:
- Treat "1." specially -> "1:" (no "1.:").
- Shows a progress bar.
- Uses multiprocessing to compute edits in parallel, then applies edits on main thread.

Usage:
    python add_colon_headings_parallel.py input.docx output.docx [--style "Heading 1"]

Requires:
    pip install python-docx tqdm
"""
from __future__ import annotations
import re
import sys
import os
from typing import List, Optional, Tuple, Dict, Any
from docx import Document
from tqdm import tqdm
from concurrent.futures import ProcessPoolExecutor, as_completed

# Matches leading numeric token like "1", "1.2", "1.2.3" (digits separated by dots).
NUM_RE = re.compile(r'^\s*(\d+(?:\.\d+)*)')

def find_run_index_and_offset(run_texts: List[str], char_index: int) -> Tuple[int,int]:
    """
    Given list of run_texts and a global char_index, return (run_idx, offset_in_run).
    If char_index == total_length, returns (last_run_idx, len(last_run_text)) to allow insertion at end.
    """
    cum = 0
    for i, t in enumerate(run_texts):
        next_cum = cum + len(t)
        if char_index <= next_cum:
            return i, char_index - cum
        cum = next_cum
    # If char_index beyond end, place at end of last run (or 0,0 if no runs)
    if run_texts:
        return len(run_texts) - 1, len(run_texts[-1])
    return 0, 0

def modify_runs_for_insertion_and_optional_deletion(run_texts: List[str],
                                                    insert_pos: int,
                                                    delete_pos: Optional[int] = None,
                                                    insert_text: str = ':') -> List[str]:
    """
    Insert insert_text at insert_pos (global index) and optionally delete one character at delete_pos (global index).
    Returns modified run_texts list (same number of runs).
    """
    # Work on a copy
    rt = list(run_texts)

    # If delete_pos provided and delete_pos < insert_pos, adjust insert_pos after deletion
    if delete_pos is not None and delete_pos < insert_pos:
        insert_pos -= 1

    # Deletion (single character) first — if requested
    if delete_pos is not None:
        ridx, roff = find_run_index_and_offset(rt, delete_pos)
        s = rt[ridx]
        # Safety: if roff is at end, nothing to delete
        if 0 <= roff < len(s):
            rt[ridx] = s[:roff] + s[roff+1:]
        # if roff == len(s) (deleting just past end), no-op

    # Now insertion
    ridx, roff = find_run_index_and_offset(rt, insert_pos)
    s = rt[ridx]
    rt[ridx] = s[:roff] + insert_text + s[roff:]

    return rt

def compute_modified_run_texts(run_texts: List[str], full_text: str) -> Optional[List[str]]:
    """
    Given run_texts and paragraph full_text, return modified run_texts if change needed,
    else return None.

    This function contains the logic for numeric detection and replacement/insertion.
    It's pure-data and safe to run in worker processes.
    """
    m = NUM_RE.match(full_text)
    if not m:
        return None  # doesn't start with a number

    num_end = m.end()  # index right after the numeric token in full_text

    # If there's already a colon immediately after the numeric token -> no change
    if num_end < len(full_text) and full_text[num_end] == ':':
        return None

    # If there's a dot '.' immediately after the number, we will replace that dot with ':'.
    delete_pos = None
    if num_end < len(full_text) and full_text[num_end] == '.':
        delete_pos = num_end
        insert_pos = num_end  # will insert colon at same logical index (helper will adjust)
    else:
        insert_pos = num_end

    # If no runs, fallback: simulate by returning a single-run list
    if not run_texts:
        # build new single run text
        if delete_pos is not None:
            # delete the dot and insert colon
            new_text = full_text[:insert_pos] + ':' + full_text[insert_pos+1:]
        else:
            new_text = full_text[:insert_pos] + ':' + full_text[insert_pos:]
        return [new_text]

    new_run_texts = modify_runs_for_insertion_and_optional_deletion(
        run_texts, insert_pos=insert_pos, delete_pos=delete_pos, insert_text=':'
    )

    return new_run_texts

def worker_task(item: Dict[str,Any]) -> Tuple[int, Optional[List[str]]]:
    """
    Worker receives a dict:
      { 'index': int, 'style_name': str, 'full_text': str, 'run_texts': List[str] }

    Returns (index, new_run_texts or None).
    """
    idx = item['index']
    run_texts = item['run_texts']
    full_text = item['full_text']
    try:
        new_runs = compute_modified_run_texts(run_texts, full_text)
        return idx, new_runs
    except Exception:
        # On any unexpected failure, return None (no change).
        return idx, None

def process_docx_parallel(in_path: str, out_path: str, target_style: str = 'Heading 1') -> int:
    doc = Document(in_path)

    # Build serializable task list: only paragraphs that match style (so less tasks)
    tasks: List[Dict[str,Any]] = []
    para_indices: List[int] = []
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style is not None else ''
        if style_name == target_style:
            # collect run texts and paragraph full text
            tasks.append({
                'index': i,
                'style_name': style_name,
                'full_text': para.text,
                'run_texts': [r.text for r in para.runs]
            })
            para_indices.append(i)

    if not tasks:
        print("No paragraphs with style", target_style, "found.")
        doc.save(out_path)
        return 0

    # Decide worker count
    max_workers = min(32, (os.cpu_count() or 1))
    results_by_index: Dict[int, Optional[List[str]]] = {}

    # Use ProcessPoolExecutor to parallelize compute-heavy work
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(worker_task, task): task['index'] for task in tasks}
        for fut in tqdm(as_completed(futures), total=len(futures), desc="Processing headings", unit="para"):
            idx = futures[fut]
            try:
                res_idx, new_runs = fut.result()
                results_by_index[res_idx] = new_runs
            except Exception as e:
                # failed worker -> treat as no change
                results_by_index[idx] = None

    # Apply changes back to the Document on the main thread
    changed = 0
    for task in tasks:
        i = task['index']
        new_runs = results_by_index.get(i)
        if not new_runs:
            continue  # no change or failure
        para = doc.paragraphs[i]
        # If paragraph has runs, assign back to existing runs (one-to-one)
        if para.runs:
            # Safety: if number of runs mismatches, we'll try to map as many as possible,
            # and if new_runs longer, append to last run; if shorter, set remaining runs to ''
            for j, run in enumerate(para.runs):
                if j < len(new_runs):
                    run.text = new_runs[j]
                else:
                    run.text = ''
            # If new_runs has extra pieces beyond number of runs, append them to last run
            if len(new_runs) > len(para.runs):
                para.runs[-1].text += ''.join(new_runs[len(para.runs):])
        else:
            # No runs in paragraph: create a single run with the text
            r = para.add_run(new_runs[0] if new_runs else '')
            # if more fragments exist (shouldn't), append them
            if len(new_runs) > 1:
                r.text += ''.join(new_runs[1:])

        changed += 1

    doc.save(out_path)
    return changed


if __name__ == "__main__":
    file_root = "C:\\DATA\\Novels\\Daily News - The Prophet of Hogwarts"
    input_file = os.path.join(file_root, "input.docx")
    output_file = os.path.join(file_root, "output.docx")
    target_style = "Heading 1"

    print("Scanning and processing (parallel). This may use multiple CPU cores.")
    modified = process_docx_parallel(input_file, output_file, target_style)
    print(f"Done. Modified {modified} paragraphs. Saved: {output_file}")
