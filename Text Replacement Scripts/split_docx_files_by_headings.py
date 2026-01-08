from docx import Document
from docx.shared import Pt, RGBColor
import os
import math
from concurrent.futures import ProcessPoolExecutor, as_completed
from tqdm import tqdm  # progress bar

def extract_paragraph_data(paragraph):
    """Extract text and formatting info from a paragraph."""
    runs_data = []
    for run in paragraph.runs:
        fmt = run.font
        run_info = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "size": fmt.size.pt if fmt.size else None,
            "color": str(fmt.color.rgb) if fmt.color and fmt.color.rgb else None,
            "name": fmt.name,
        }
        runs_data.append(run_info)
    return {
        "runs": runs_data,
        "style": paragraph.style.name
    }

def extract_sections(doc, heading_style):
    """Extract sections grouped by headings."""
    sections = []
    current = None
    for para in doc.paragraphs:
        if para.style.name == heading_style:
            if current is not None:
                sections.append(current)
            current = [extract_paragraph_data(para)]
        else:
            if current is not None:
                current.append(extract_paragraph_data(para))
    if current:
        sections.append(current)
    return sections

def write_docx(group, output_path):
    """Worker: writes a DOCX file from serialized section data."""
    part = Document()
    available_styles = {s.name for s in part.styles}

    for section in group:
        for para_data in section:
            style_name = para_data["style"] if para_data["style"] in available_styles else "Normal"
            para = part.add_paragraph(style=style_name)
            for run_info in para_data["runs"]:
                run = para.add_run(run_info["text"])
                if run_info["bold"]:
                    run.bold = True
                if run_info["italic"]:
                    run.italic = True
                if run_info["underline"]:
                    run.underline = True
                if run_info["size"]:
                    run.font.size = Pt(run_info["size"])
                if run_info["color"]:
                    run.font.color.rgb = RGBColor.from_string(run_info["color"])
                if run_info["name"]:
                    run.font.name = run_info["name"]

    part.save(output_path)
    return output_path

def split_docx_by_headings_parallel(input_path, output_dir, heading_style="Heading 1", num_files=None, max_workers=4):
    """Split DOCX by headings, preserving formatting, with parallel progress bar."""
    doc = Document(input_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Step 1: Extract and serialize sections
    sections = extract_sections(doc, heading_style)

    # Step 2: Group into chunks for desired num_files
    if num_files is None or num_files >= len(sections):
        split_groups = [[s] for s in sections]
    else:
        split_groups = []
        per_file = math.ceil(len(sections) / num_files)
        for i in range(0, len(sections), per_file):
            split_groups.append(sections[i:i + per_file])

    total_parts = len(split_groups)
    print(f"Starting parallel split using {min(max_workers, total_parts)} processes...")

    # Step 3: Parallel writing with progress bar
    futures = []
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        for idx, group in enumerate(split_groups, start=1):
            output_path = os.path.join(output_dir, f"part_{idx}.docx")
            futures.append(executor.submit(write_docx, group, output_path))

        # tqdm progress bar
        with tqdm(total=total_parts, desc="Writing DOCX files", ncols=80) as pbar:
            for future in as_completed(futures):
                future.result()  # wait for completion
                pbar.update(1)

    print(f"✅ Split completed. {total_parts} files saved to {output_dir}")

# Example usage
if __name__ == "__main__":
    file_root = "C:\\DATA\\Novels\\I Really am an Inside Player"
    input_file = os.path.join(file_root, "input.docx")
    output_directory = os.path.join(file_root, "output_parts")

    split_docx_by_headings_parallel(
        input_file,
        output_directory,
        heading_style="Heading 1",
        num_files=3,
        max_workers=16
    )