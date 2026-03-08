import docx
import re
import os
import unicodedata
from concurrent.futures import ProcessPoolExecutor
from tqdm import tqdm

# --- EPUB Imports ---
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup, NavigableString

# --- Mappings ---
PUNCT_REPLACEMENTS = [
    # 1. Curly quotes -> straight
    (r'[\u201C\u201D\u201F\u300C\u300D\u300E\u300F]', '"'), 
    (r'[\u2018\u2019\u201B]', "'"),
    # 2. Ellipsis
    (r'\u2026', '...'),
    # 3. CJK -> ASCII
    (r'，', ','), (r'。', '.'), (r'：', ':'), (r'；', ';'),
    (r'！', '!'), (r'？', '?'), (r'（', '('), (r'）', ')'),
    (r'【', '['), (r'】', ']'), (r'、', ','), 
    # 4. Dashes
    (r'——', '—'), 
    # 5. Cleanup
    (r'\u00A0', ' '), 
]

def clean_text_worker(text):
    """
    Parallel Worker: Cleans text INSIDE a single run/node.
    """
    if not text: return ""

    # 1. Normalize Unicode Spaces
    text = unicodedata.normalize('NFKC', text)
    text = re.sub(r'[\u00A0\u2000-\u200B\u3000]', ' ', text)

    # 2. Mappings
    for pattern, replacement in PUNCT_REPLACEMENTS:
        text = re.sub(pattern, replacement, text)

    # 3. Numbers/Currency/Contractions
    text = re.sub(r'(\d)\s*\.\s*(\d)', r'\1.\2', text) # Fix decimals 3 . 5 -> 3.5
    text = re.sub(r'(\d)\s*,\s*(\d{3})', r'\1,\2', text)
    text = re.sub(r'(\d)\s+%', r'\1%', text)
    text = re.sub(r'(\$|€|£|¥)\s+(\d)', r'\1\2', text)
    text = re.sub(r'([a-zA-Z])\s+\'', r'\1\'', text)
    text = re.sub(r'\'\s+(s|t|m|re|ll|ve|d)(?!\w)', r"'\1", text, flags=re.IGNORECASE)
    text = re.sub(r'\b(i)\b', 'I', text)

    # 4. Unstick Quotes
    text = re.sub(r'(?<=[a-zA-Z])"(?=[A-Z])', ' "', text)
    
    # 5. Internal Spacing Fixes
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'"\s+(?=\S)', '"', text)       
    text = re.sub(r'(?<=[a-zA-Z0-9.,!?])\s+"', '"', text)  
    text = re.sub(r'\(\s+', '(', text)
    text = re.sub(r'\s+\)', ')', text)

    # 6. Punctuation Spacing & NEW FIX: Space after dot before Sentence
    # Ensures "word.Next" becomes "word. Next"
    # Look for (.!?) followed immediately by a Capital letter
    text = re.sub(r'([.!?])(?=[A-Z])', r'\1 ', text)
    
    # Standard cleanup for other punctuation followed by words
    text = re.sub(r'([,;:])(?=[a-zA-Z])', r'\1 ', text)
    text = re.sub(r'([.,;:!?])(?=")', r'\1 ', text)

    # 7. Deduplicate Punctuation
    text = re.sub(r'([,:;])\1+', r'\1', text)
    text = re.sub(r'([!?])\1+', r'\1', text)
    text = re.sub(r'\.{4,}', '...', text)

    # 8. Sentence Capitalization
    if len(text) > 0 and text[0].islower():
        # Only capitalize if it looks like the start of a sentence (simple heuristic)
        text = text[0].upper() + text[1:]
    text = re.sub(r'([.!?]\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text)

    return text

# ==========================================
# DOCX HANDLING
# ==========================================

def fix_split_run_boundaries(doc):
    """
    Advanced Zipper: Fixes spacing errors that span across TWO runs.
    Updated to handle the "Space after Dot" rule across boundaries.
    """
    print("Stitching broken run boundaries...")
    
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    count = 0
    CLOSING_PUNCT = {'.', ',', '!', '?', ';', ':'}
    
    for para in tqdm(all_paras, desc="Stitching Boundaries"):
        runs = para.runs
        if len(runs) < 2: continue
            
        for i in range(len(runs) - 1):
            curr_run = runs[i]
            next_run = runs[i+1]
            
            if not curr_run.text or not next_run.text: continue

            c_txt = curr_run.text
            n_txt = next_run.text
            c_stripped = c_txt.strip()
            n_stripped = n_txt.strip()

            if not c_stripped or not n_stripped: continue

            # --- 1. EXISTING SPACING LOGIC ---
            if (c_stripped.endswith('"') or c_stripped.endswith('(')) and c_txt.endswith(' '):
                 curr_run.text = c_txt.rstrip()
                 count += 1
            elif (c_txt.endswith('"') or c_txt.endswith('(')) and n_txt.startswith(' '):
                 next_run.text = n_txt.lstrip()
                 count += 1
            elif c_txt.endswith(' ') and (n_txt.startswith('"') or n_txt.startswith(')') or n_txt[0] in CLOSING_PUNCT):
                 curr_run.text = c_txt.rstrip()
                 count += 1
            elif n_txt.startswith(' ') and (n_stripped.startswith('"') or n_stripped.startswith(')') or n_stripped[0] in CLOSING_PUNCT):
                 next_run.text = n_txt.lstrip()
                 count += 1

            # --- 2. NEW LOGIC: Dot at end of Run A, Capital at start of Run B ---
            # Case: A="end." B="Start" -> Should be "end." B=" Start"
            if c_stripped.endswith('.') and n_txt[0].isupper() and not c_txt.endswith(' ') and not n_txt.startswith(' '):
                # We need to insert a space. Safer to add to the start of B.
                next_run.text = ' ' + n_txt
                count += 1

    print(f"Stitched {count} split boundaries.")

def clean_paragraph_ends(doc):
    """Removes leading/trailing whitespace from paragraphs."""
    print("Cleaning paragraph start/end...")
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)
    
    for para in all_paras:
        if not para.runs: continue
        # Fix Leading
        for run in para.runs:
            if run.text:
                if run.text.startswith(' '): run.text = run.text.lstrip()
                break 
        # Fix Trailing
        for run in reversed(para.runs):
            if run.text:
                if run.text.endswith(' '): run.text = run.text.rstrip()
                break

def collect_docx_runs(doc):
    all_runs = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text: all_runs.append(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text: all_runs.append(run)
    return all_runs

def process_docx(input_path, output_path, max_workers):
    print(f"Processing DOCX: {input_path}")
    doc = docx.Document(input_path)

    # 1. Parallel Regex Cleaning
    print("Step 1: Parallel Regex Cleaning...")
    all_runs = collect_docx_runs(doc)
    raw_texts = [run.text for run in all_runs]
    
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        cleaned_texts = list(tqdm(executor.map(clean_text_worker, raw_texts), total=len(raw_texts), unit="run"))

    for i, run in enumerate(all_runs):
        if run.text != cleaned_texts[i]:
            run.text = cleaned_texts[i]

    # 2. Stitch broken boundaries
    print("Step 2: Stitching Boundaries...")
    fix_split_run_boundaries(doc)
    
    # 3. Clean Paragraph Start/Ends
    print("Step 3: Trimming Paragraphs...")
    clean_paragraph_ends(doc)

    print(f"Saving to {output_path}...")
    doc.save(output_path)

# ==========================================
# EPUB HANDLING
# ==========================================

def process_epub(input_path, output_path, max_workers):
    print(f"Processing EPUB: {input_path}")
    try:
        book = epub.read_epub(input_path)
    except Exception as e:
        print(f"Error reading EPUB: {e}")
        return

    # Collect all HTML items (Chapters)
    html_items = [item for item in book.get_items() if item.get_type() == ebooklib.ITEM_DOCUMENT]
    print(f"Found {len(html_items)} text documents in EPUB.")

    # We process items sequentially, but the text *inside* items in parallel
    for item in tqdm(html_items, desc="Processing Chapters"):
        content = item.get_content()
        soup = BeautifulSoup(content, 'html.parser')
        
        # 1. Extract all text nodes (NavigableStrings)
        # We filter out scripts and styles to avoid breaking code
        text_nodes = []
        for node in soup.find_all(string=True):
            if isinstance(node, NavigableString):
                parent = node.parent
                if parent.name not in ['script', 'style', 'code', 'pre']:
                    text_nodes.append(node)
        
        if not text_nodes:
            continue

        raw_texts = [str(node) for node in text_nodes]

        # 2. Parallel Clean
        with ProcessPoolExecutor(max_workers=max_workers) as executor:
            cleaned_texts = list(executor.map(clean_text_worker, raw_texts))

        # 3. Replace text in Soup
        for i, node in enumerate(text_nodes):
            if raw_texts[i] != cleaned_texts[i]:
                node.replace_with(cleaned_texts[i])

        # 4. Save modified HTML back to item
        # Use utf-8 to ensure special chars are kept
        item.set_content(soup.encode(formatter="html"))

    print(f"Saving EPUB to {output_path}...")
    epub.write_epub(output_path, book)

# ==========================================
# MAIN
# ==========================================

if __name__ == "__main__":
    # Settings
    FILE_ROOT = "C:\\DATA\\Novels\\Football Manager - Running a Rip-off club"
    INPUT_FILENAME = "Football Manager - Running a Rip-off club.epub" # .docx or .epub
    OUTPUT_FILENAME = "output.epub" # .docx or .epub
    
    INPUT_FILE = os.path.join(FILE_ROOT, INPUT_FILENAME)
    OUTPUT_FILE = os.path.join(FILE_ROOT, OUTPUT_FILENAME)
    
    MAX_WORKERS = os.cpu_count()

    if not os.path.exists(INPUT_FILE):
        print(f"File not found: {INPUT_FILE}")
    else:
        ext = os.path.splitext(INPUT_FILE)[1].lower()
        
        if ext == ".docx":
            process_docx(INPUT_FILE, OUTPUT_FILE, MAX_WORKERS)
            print("Docx Done!")
        elif ext == ".epub":
            process_epub(INPUT_FILE, OUTPUT_FILE, MAX_WORKERS)
            print("Epub Done!")
        else:
            print(f"Unsupported file format: {ext}")