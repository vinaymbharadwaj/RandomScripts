import os
import json
import ebooklib
import ollama
from tqdm import tqdm
import docx
from ebooklib import epub
from bs4 import BeautifulSoup, NavigableString

# --- Configuration & Setup ---
MODEL_NAME = 'qwen3:4b'

def load_glossary(file_path: str) -> dict:
    """Loads glossary from a text file formatted as Chinese=English."""
    glossary = {}
    if not file_path or not os.path.exists(file_path):
        return glossary
        
    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if '=' in line:
                zh, en = line.split('=', 1)
                glossary[zh.strip()] = en.strip()
    return glossary

def get_system_instruction(glossary: dict) -> str:
    """Builds the strict instruction prompt, dynamically injecting the glossary if it exists."""
    base_instruction = """You are a strict, literal translator for Chinese web novels. 
Translate the provided Chinese text into English.

CRITICAL RULES:
1. BRACKET PRESERVATION: Sentences, dialogue, or notes enclosed in brackets [] or parentheses () MUST remain on their own separate lines. You are strictly forbidden from merging them into surrounding text.
2. LINE BREAKS: Maintain all original paragraph breaks.
3. OUTPUT ONLY: Return only the translated English text. Do not include notes, conversational filler, or explanations.
"""
    if glossary:
        glossary_rule = f"\n4. EXACT GLOSSARY: You must translate the following terms exactly as specified: {json.dumps(glossary, ensure_ascii=False)}. Do not deviate.\n"
        base_instruction += glossary_rule
        
    return base_instruction

def translate_chunk(text: str, system_prompt: str) -> str:
    """Sends a chunk of text to the local GPU."""
    if not text.strip():
        return text
        
    try:
        response = ollama.chat(
            model=MODEL_NAME, 
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': text}
            ],
            options={"temperature": 0.1}
        )
        return response['message']['content'].strip()
    except Exception as e:
        print(f"\nTranslation Error: {e}")
        return text

# --- File Handlers ---

def process_txt(input_path: str, output_path: str, system_prompt: str):
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    translated_lines = []
    with tqdm(total=len(lines), desc="Translating TXT", unit="line") as pbar:
        for line in lines:
            if line.strip():
                translated_lines.append(translate_chunk(line, system_prompt) + "\n")
            else:
                translated_lines.append("\n")
            pbar.update(1)
            
    print("Generating TXT output...")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.writelines(translated_lines)

def process_docx(input_path: str, output_path: str, system_prompt: str):
    print("Reading DOCX input...")
    doc = docx.Document(input_path)
    
    with tqdm(total=len(doc.paragraphs), desc="Translating DOCX", unit="para") as pbar:
        for para in doc.paragraphs:
            original_text = para.text.strip()
            if original_text:
                # 1. Save paragraph structural styles (Headings, Title, Normal, etc.)
                p_style = para.style
                p_alignment = para.alignment
                
                # 2. Quick heuristic for inline styles (Bold/Italic)
                # We check the first non-empty run. If the original text started bold/italic, 
                # we apply that to the entire translated line. This avoids slow run-by-run translation.
                is_bold = False
                is_italic = False
                for run in para.runs:
                    if run.text.strip():
                        is_bold = run.bold
                        is_italic = run.italic
                        break 
                        
                # 3. Translate the paragraph text all at once (Max speed, preserves AI context)
                translated_text = translate_chunk(original_text, system_prompt)
                
                # 4. Clear the old Chinese text and rebuild the paragraph with the saved styles
                para.clear()
                para.style = p_style
                if p_alignment is not None:
                    para.alignment = p_alignment
                    
                new_run = para.add_run(translated_text)
                
                # Re-apply inline formatting if it existed
                if is_bold:
                    new_run.bold = True
                if is_italic:
                    new_run.italic = True
                
            pbar.update(1)
            
    print("Generating DOCX output...")
    doc.save(output_path)

def process_epub(input_path: str, output_path: str, system_prompt: str):
    print("Reading EPUB input...")
    book = epub.read_epub(input_path)
    
    items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
    
    with tqdm(total=len(items), desc="Translating EPUB chapters", unit="ch") as pbar:
        for item in items:
            soup = BeautifulSoup(item.get_body_content(), 'html.parser')
            
            # Find all text nodes. Translating node-by-node preserves HTML structure.
            text_nodes = soup.find_all(string=True)
            text_nodes = [n for n in text_nodes if n.strip() and not isinstance(n, NavigableString) == False]
            
            if text_nodes:
                # Sub-progress bar for paragraphs within a chapter
                for node in tqdm(text_nodes, desc=f"  -> Chapter nodes", leave=False):
                    original_text = str(node)
                    if original_text.strip():
                        translated = translate_chunk(original_text, system_prompt)
                        node.replace_with(translated)
                        
            item.set_content(str(soup).encode('utf-8'))
            pbar.update(1)
            
    print("Generating EPUB output...")
    epub.write_epub(output_path, book)

# --- Main Execution ---

def main():
    # 1. Define your inputs here
    file_root = "C:\\DATA\\Novels\\Oregairu - This isn't right, Why is it Conan"
    input_file = os.path.join(file_root, "input.docx")  # can be input.txt or input.docx or input.epub
    output_file = os.path.join(file_root, "output.docx")  # .txt or .docx
    glossary_file = None  # Set to None if you don't want to use a glossary
    
    if not os.path.exists(input_file):
        print(f"Error: Input file {input_file} not found.")
        return

    # 2. Load Glossary and Prompt
    print("Initializing...")
    glossary = load_glossary(glossary_file)
    if glossary:
        print(f"Loaded {len(glossary)} glossary terms.")
    else:
        print("Running without glossary.")
        
    system_prompt = get_system_instruction(glossary)
    
    # 3. Route to correct processor based on extension
    ext = input_file.lower().split('.')[-1]
    
    if ext == 'txt':
        process_txt(input_file, output_file, system_prompt)
    elif ext == 'docx':
        process_docx(input_file, output_file, system_prompt)
    elif ext == 'epub':
        process_epub(input_file, output_file, system_prompt)
    else:
        print(f"Unsupported file format: {ext}")
        return
        
    # 4. Cleanup
    print("\nTranslation complete. Flushing VRAM...")
    ollama.chat(model=MODEL_NAME, keep_alive=0)
    print("GPU resources released.")

if __name__ == "__main__":
    main()