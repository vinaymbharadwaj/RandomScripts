from docx import Document
import os
import math

def split_docx_by_headings(input_path, output_dir, heading_style="Heading 1", num_files=None):
    """
    Splits a DOCX file into multiple DOCX files based on headings.
    User can specify how many output files to generate.

    :param input_path: Path to the input .docx file
    :param output_dir: Directory to save the split .docx files
    :param heading_style: Which heading style to split on (e.g., "Heading 1")
    :param num_files: Desired number of split files (if None, split by every heading)
    """
    # Load the document
    doc = Document(input_path)

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Step 1: Collect sections based on headings
    sections = []
    current = None

    for para in doc.paragraphs:
        if para.style.name == heading_style:
            if current is not None:
                sections.append(current)
            current = [(para.text, para.style)]
        else:
            if current is not None:
                current.append((para.text, para.style))

    if current:
        sections.append(current)

    # Step 2: If user wants fixed number of files, regroup sections
    if num_files is None or num_files >= len(sections):
        split_groups = [[s] for s in sections]  # one file per heading
    else:
        # group sections into num_files
        split_groups = []
        per_file = math.ceil(len(sections) / num_files)
        for i in range(0, len(sections), per_file):
            split_groups.append(sections[i:i+per_file])

    # Step 3: Write files
    for idx, group in enumerate(split_groups, start=1):
        part = Document()
        for section in group:
            for text, style in section:
                part.add_paragraph(text, style=style)
        output_path = os.path.join(output_dir, f"part_{idx}.docx")
        part.save(output_path)

    print(f"Split completed. {len(split_groups)} files saved to {output_dir}")

# Example usage:
file_root = "C:\\DATA\\Novels\\I Took in the Pokémon"
input_file = os.path.join(file_root, "input.docx")
output_directory = os.path.join(file_root, "output_parts")
split_docx_by_headings(input_file, output_directory, heading_style="Heading 1", num_files=3)